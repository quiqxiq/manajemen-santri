# -*- coding: utf-8 -*-
"""
Build BAB IV & BAB V (skripsi SIPADES) as .docx matching PROPOSAL-SKRIPSI.docx
formatting exactly: A4, margins 4/3/3/3 cm, 1.5 line spacing (line=360),
TNR, Heading1 (BAB) centered bold, Heading2/3 numbered, justified body with
firstLine indent, centered figure captions, shaded table headers, PAGE footer.
"""
import os, sys
from PIL import Image
from docx import Document
from docx.shared import Twips, Pt, Inches, Emu, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.abspath(os.path.join(BASE, "..", "screenshots"))
OUT = os.path.abspath(os.path.join(BASE, "..", "screenshots", "BAB-IV-V-SIPADES.docx"))

# ---------------- low-level formatting helpers ----------------

def _child(parent, tag):
    el = parent.find(qn(tag))
    if el is None:
        el = OxmlElement(tag)
        parent.append(el)
    return el

def set_spacing(p, before=None, after=0, line=360, rule="auto"):
    pPr = p._p.get_or_add_pPr()
    sp = _child(pPr, "w:spacing")
    if before is not None:
        sp.set(qn("w:before"), str(before))
    if after is not None:
        sp.set(qn("w:after"), str(after))
    if line is not None:
        sp.set(qn("w:line"), str(line))
        sp.set(qn("w:lineRule"), rule)

def set_indent(p, firstLine=None, left=None, hanging=None):
    pPr = p._p.get_or_add_pPr()
    ind = _child(pPr, "w:ind")
    if firstLine is not None:
        ind.set(qn("w:firstLine"), str(firstLine))
    if left is not None:
        ind.set(qn("w:left"), str(left))
    if hanging is not None:
        ind.set(qn("w:hanging"), str(hanging))

def style_run(r, pt=12, bold=False, italic=False, color=None, name="Times New Roman"):
    r.font.name = name
    rpr = r._r.get_or_add_rPr()
    rf = _child(rpr, "w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(a), name)
    r.font.size = Pt(pt)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = RGBColor.from_string(color)

# ---------------- page setup ----------------

def setup_page(doc):
    """A4, margins: top/right/bottom=3cm, left=4cm; footer distance 1.27cm."""
    sec = doc.sections[0]
    sec.page_width  = Twips(11907)
    sec.page_height = Twips(16840)
    sec.top_margin    = Twips(1701)
    sec.right_margin  = Twips(1701)
    sec.bottom_margin = Twips(1701)
    sec.left_margin   = Twips(2268)
    sec.footer_distance = Twips(720)
    sec.header_distance = Twips(720)

# ---------------- footer (centered PAGE field) ----------------

def add_page_footer(doc):
    sec = doc.sections[0]
    sec.different_first_page_header_footer = False
    footer = sec.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)

    sdt = OxmlElement("w:sdt")
    sdtPr = OxmlElement("w:sdtPr")
    sdtId = OxmlElement("w:id"); sdtId.set(qn("w:val"), "-264543138")
    sdtPr.append(sdtId)
    dpo = OxmlElement("w:docPartObj")
    dpg = OxmlElement("w:docPartGallery"); dpg.set(qn("w:val"), "Page Numbers (Bottom of Page)")
    dpu = OxmlElement("w:docPartUnique")
    dpo.append(dpg); dpo.append(dpu); sdtPr.append(dpo)
    sdt.append(sdtPr)
    sdtEnd = OxmlElement("w:sdtEndPr")
    rpr_end = OxmlElement("w:rPr"); np_end = OxmlElement("w:noProof"); rpr_end.append(np_end)
    sdtEnd.append(rpr_end); sdt.append(sdtEnd)
    sdtContent = OxmlElement("w:sdtContent")

    p_el = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle"); pStyle.set(qn("w:val"), "Footer"); pPr.append(pStyle)
    jc = OxmlElement("w:jc"); jc.set(qn("w:val"), "center"); pPr.append(jc)
    p_el.append(pPr)
    def fld_run(ftype):
        r = OxmlElement("w:r"); fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), ftype); r.append(fc); return r
    p_el.append(fld_run("begin"))
    r_instr = OxmlElement("w:r"); it = OxmlElement("w:instrText")
    it.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    it.text = " PAGE   \\* MERGEFORMAT "; r_instr.append(it); p_el.append(r_instr)
    p_el.append(fld_run("separate"))
    r_val = OxmlElement("w:r"); rpr_v = OxmlElement("w:rPr"); np_v = OxmlElement("w:noProof")
    rpr_v.append(np_v); r_val.append(rpr_v)
    t_v = OxmlElement("w:t"); t_v.text = "1"; r_val.append(t_v); p_el.append(r_val)
    p_el.append(fld_run("end"))
    sdtContent.append(p_el); sdt.append(sdtContent)
    footer._element.append(sdt)

# ---------------- paragraph factories ----------------

def add_bab_heading(doc, roman, title):
    """Heading1: 'BAB IV\nHASIL DAN PEMBAHASAN' centered, TNR 12pt bold."""
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    pPr = p._p.get_or_add_pPr()
    sp = _child(pPr, "w:spacing")
    sp.set(qn("w:before"), "0"); sp.set(qn("w:after"), "120")
    sp.set(qn("w:line"), "360"); sp.set(qn("w:lineRule"), "auto")
    jc = _child(pPr, "w:jc"); jc.set(qn("w:val"), "center")
    r1 = p.add_run("BAB " + roman); style_run(r1, pt=12, bold=True)
    r_br = p.add_run(); r_br._r.append(OxmlElement("w:br"))
    r2 = p.add_run(title); style_run(r2, pt=12, bold=True)
    return p

def add_h2(doc, num, title):
    """Heading2: '4.1  Hasil Penelitian' — TNR 12pt bold, before=240, after=60."""
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    pPr = p._p.get_or_add_pPr()
    sp = _child(pPr, "w:spacing")
    sp.set(qn("w:before"), "240"); sp.set(qn("w:after"), "60")
    sp.set(qn("w:line"), "360"); sp.set(qn("w:lineRule"), "auto")
    r = p.add_run(num + "  " + title); style_run(r, pt=12, bold=True)
    return p

def add_h3(doc, num, title):
    """Heading3: '4.1.1  Lingkungan Implementasi' — TNR 12pt bold, before=180, after=60."""
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 3"]
    pPr = p._p.get_or_add_pPr()
    sp = _child(pPr, "w:spacing")
    sp.set(qn("w:before"), "180"); sp.set(qn("w:after"), "60")
    sp.set(qn("w:line"), "360"); sp.set(qn("w:lineRule"), "auto")
    r = p.add_run(num + "  " + title); style_run(r, pt=12, bold=True)
    return p

def add_body(doc, text, bold_prefix=None, italic=False):
    """Normal body: justified, firstLine=426 (0.75 cm), TNR 12pt, line=360."""
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    set_spacing(p, before=0, after=0, line=360)
    set_indent(p, firstLine=426)
    p.alignment = AL.JUSTIFY
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        style_run(r_pre, pt=12, bold=True, italic=italic)
    r = p.add_run(text)
    style_run(r, pt=12, italic=italic)
    return p

def add_caption(doc, text):
    """Caption: centered, TNR 11pt bold label, before=120 after=120 line=360."""
    p = doc.add_paragraph()
    p.style = doc.styles["Caption"]
    pPr = p._p.get_or_add_pPr()
    sp = _child(pPr, "w:spacing")
    sp.set(qn("w:before"), "120"); sp.set(qn("w:after"), "120")
    sp.set(qn("w:line"), "360"); sp.set(qn("w:lineRule"), "auto")
    jc = _child(pPr, "w:jc"); jc.set(qn("w:val"), "center")
    ind = _child(pPr, "w:ind"); ind.set(qn("w:firstLine"), "0")
    
    # Split caption for bold label if starts with Gambar or Tabel
    if text.startswith("Gambar ") or text.startswith("Tabel "):
        parts = text.split("  ", 1)
        if len(parts) == 2:
            r1 = p.add_run(parts[0] + "  "); style_run(r1, pt=11, bold=True, color="000000")
            r2 = p.add_run(parts[1]); style_run(r2, pt=11, bold=False, color="000000")
        else:
            r = p.add_run(text); style_run(r, pt=11, bold=False, color="000000")
    else:
        r = p.add_run(text); style_run(r, pt=11, bold=False, color="000000")
    return p

# ---------------- image insertion ----------------

IMG_COUNTER = {"n": 0}

def add_figure(doc, filename, caption_text):
    IMG_COUNTER["n"] += 1
    num_str = f"Gambar 4.{IMG_COUNTER['n']}"
    full_caption = f"{num_str}  {caption_text}"
    path = os.path.join(IMG_DIR, filename)
    try:
        with Image.open(path) as im:
            w_px, h_px = im.size
    except Exception:
        w_px, h_px = 1280, 720
    
    # Max width: 14 cm (~5.51 in)
    max_w_emu = int(5.51 * 914400)
    max_h_emu = int(3.5 * 914400)
    ratio = w_px / h_px
    w_emu = max_w_emu
    h_emu = int(w_emu / ratio)
    if h_emu > max_h_emu:
        h_emu = max_h_emu
        w_emu = int(h_emu * ratio)
        
    p = doc.add_paragraph()
    set_spacing(p, before=180, after=60, line=360)
    p.alignment = AL.CENTER
    run = p.add_run()
    run.add_picture(path, width=Emu(w_emu), height=Emu(h_emu))
    add_caption(doc, full_caption)
    return p

# ---------------- black-box test table ----------------

def add_bb_table(doc, tbl_num, caption_text, rows):
    """rows = list of (no, skenario, input, expected, result, status)"""
    add_caption(doc, f"Tabel {tbl_num}  {caption_text}")
    tbl = doc.add_table(rows=1 + len(rows), cols=6)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Column widths (Total ~ 14 cm)
    # No: 0.8 cm, Skenario: 2.8 cm, Input: 2.8 cm, Expected: 3.3 cm, Result: 3.3 cm, Status: 1.0 cm
    widths = [Twips(450), Twips(1580), Twips(1580), Twips(1860), Twips(1860), Twips(600)]
    for i, w in enumerate(widths):
        for cell in tbl.columns[i].cells:
            cell.width = w
            
    headers = ["No.", "Skenario Uji", "Data Masukan", "Hasil yang Diharapkan", "Hasil Pengujian", "Status"]
    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        set_spacing(p, before=60, after=60, line=240)
        p.alignment = AL.CENTER
        r = p.add_run(h); style_run(r, pt=10, bold=True)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "D9D9D9")
        tcPr.append(shd)
        
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            set_spacing(p, before=60, after=60, line=240)
            p.alignment = AL.CENTER if ci in (0, 5) else AL.JUSTIFY
            r = p.add_run(str(val))
            is_bold = (ci == 5)
            style_run(r, pt=9.5, bold=is_bold)
    return tbl

# ---------------- DOCUMENT GENERATOR MAIN ----------------

def generate_docx():
    doc = Document()
    setup_page(doc)
    add_page_footer(doc)
    
    print("Generating BAB IV...")
    
    # ---------------- BAB IV ----------------
    add_bab_heading(doc, "IV", "HASIL DAN PEMBAHASAN")
    
    # 4.1 Hasil Penelitian
    add_h2(doc, "4.1", "Hasil Penelitian")
    add_body(doc, "Hasil penelitian ini adalah terbangunnya Sistem Informasi Pelayanan Desa (SIPADES) Karduluk berbasis web yang dirancang untuk mengotomatisasi seluruh alur permohonan surat-menyurat di Kantor Desa Karduluk. Sistem ini memadukan konsep self-service portal untuk warga desa, mekanisme verifikasi persetujuan berjenjang (multi-level approval) dinamis, otomatisasi pengiriman notifikasi terintegrasi WhatsApp Gateway, serta pengesahan dokumen berupa Tanda Tangan Elektronik (TTE) berbasis QR Code terverifikasi.")
    
    # 4.1.1 Lingkungan Implementasi
    add_h3(doc, "4.1.1", "Lingkungan Implementasi")
    add_body(doc, "Pengembangan dan implementasi sistem SIPADES dilaksanakan pada lingkungan perangkat keras (hardware) dan perangkat lunak (software) dengan spesifikasi teknis sebagai berikut:")
    add_body(doc, "1. Spesifikasi Perangkat Keras (Hardware): Sistem dikembangkan pada unit PC/Laptop penguji berbasis prosesor Intel Core i7-12700H, memori utama (RAM) 16 GB DDR4, serta media penyimpanan Solid State Drive (SSD) NVMe 512 GB. Lingkungan peladen (server) pengujian dijalankan pada infrastruktur pengujian lokal yang siap distigmasikan ke lingkungan Cloud Server (VPS).")
    add_body(doc, "2. Spesifikasi Perangkat Lunak (Software): Sistem beroperasi pada lingkungan Sistem Operasi Microsoft Windows 11 Enterprise / Linux Ubuntu 22.04 LTS. Bahasa pemrograman utama yang digunakan adalah PHP 8.3/8.2 dengan basis web framework Laravel 13, panel manajemen administrasi Filament v5, database MySQL 8.0/SQLite 3, web server Nginx/Apache, serta browser Google Chrome untuk pengujian fungsionalitas.")
    add_body(doc, "3. Tech Stack dan Pustaka Pendukung: Arsitektur perangkat lunak memanfaatkan Blade + Livewire untuk komponen UI dinamis portal warga, Spatie Laravel Permission dan Filament Shield untuk kontrol hak akses berbasis peran (RBAC), Barryvdh Laravel DomPDF untuk pembentukan dokumen surat resmi dalam format PDF, Spatie Media Library untuk pengelolaan berkas persyaratan, Spatie Activity Log untuk pencatatan jejak audit (audit trail), serta Go-WA Gateway Service yang terhubung via HTTP REST API untuk layanan notifikasi WhatsApp.")

    # 4.1.2 Implementasi Antarmuka Portal Warga
    add_h3(doc, "4.1.2", "Implementasi Antarmuka Portal Warga")
    add_body(doc, "Antarmuka portal warga dirancang berorientasi pada kemudahan pengguna (user-friendly) dan kemudahan akses mandiri (self-service). Warga dapat mengakses halaman utama publik untuk melihat informasi umum layanan desa sebelum melakukan autentikasi.")
    
    add_figure(doc, "01_publik_landing.png", "Halaman Utama / Landing Page Publik SIPADES Karduluk")
    add_body(doc, "Halaman utama (Landing Page) menampilkan banner selamat datang, daftar informasi layanan persuratan unggulan, statistik pelayanan secara tepat waktu, serta tautan menuju portal login dan verifikasi keabsahan surat publik.")
    
    add_figure(doc, "02_portal_login.png", "Antarmuka Halaman Login Portal Warga")
    add_body(doc, "Halaman login portal warga digunakan oleh warga terdaftar untuk masuk ke dalam sistem menggunakan alamat email dan kata sandi yang telah diverifikasi.")
    
    add_figure(doc, "03_registrasi_warga.png", "Formulir Registrasi Mandiri Warga Desa Karduluk")
    add_body(doc, "Warga baru yang belum memiliki akun dapat mengisikan formulir registrasi mandiri yang mencakup Data NIK (Nomor Induk Kependudukan), Nama Lengkap sesuai KTP, Alamat Email, Nomor WhatsApp aktif, serta Kata Sandi.")
    
    add_figure(doc, "04_verifikasi_otp.png", "Antarmuka Verifikasi Kode OTP via WhatsApp")
    add_body(doc, "Untuk menjamin keabsahan data kependudukan dan kepemilikan nomor telepon yang valid, sistem mengirimkan 6 digit kode OTP (One-Time Password) secara otomatis ke WhatsApp warga saat proses registrasi atau permohonan keamanan.")
    
    add_figure(doc, "07_portal_dashboard.png", "Halaman Dashboard Utama Portal Warga")
    add_body(doc, "Setelah berhasil login, warga disambut pada Halaman Dashboard Warga yang memuat kartu informasi ringkasan pengajuan (Total Pengajuan, Dalam Proses, Disetujui, Ditolak), tombol cepat ajukan surat, serta daftar riwayat aktivitas pengajuan terbaru.")
    
    add_figure(doc, "08_portal_ajukan_surat_buat.png", "Formulir Pengajuan Permohonan Surat Pelayanan Desa")
    add_body(doc, "Warga dapat memilih jenis surat yang dibutuhkan (misal: Surat Keterangan Usaha, Surat Keterangan Dominasi, Surat Keterangan Tidak Mampu, dll.), mengisikan formulir isian khusus, serta mengunggah dokumen persyaratan seperti foto KTP atau Kartu Keluarga.")
    
    add_figure(doc, "09_portal_pengajuan_index.png", "Halaman Riwayat Pengajuan Surat Portal Warga")
    add_body(doc, "Halaman riwayat pengajuan menampilkan seluruh berkas permohonan yang pernah dibuat oleh warga beserta indikator status persetujuan secara real-time dan opsi penyaringan (filter status).")
    
    add_figure(doc, "10_portal_pengajuan_status_detail.png", "Halaman Detail Pengajuan & Visualisasi Progress Stepper Approval")
    add_body(doc, "Warga dapat melacak tahapan verifikasi surat secara rinci melalui komponen visualisasi Stepper Progress. Komponen ini menampilkan status verifikasi Level 1 (Petugas Desa), Level 2 (Sekretaris Desa), hingga Level 3 (Kepala Desa & TTE).")
    
    add_figure(doc, "11_portal_pengajuan_revisi.png", "Formulir dan Catatan Revisi Pengajuan Surat")
    add_body(doc, "Apabila petugas menemukan kekurangan atau kesalahan pada berkas persyaratan, sistem akan menampilkan catatan revisi dari verifikator dan memberikan akses bagi warga untuk memperbaiki formulir serta mengunggah ulang berkas perbaikan.")
    
    add_figure(doc, "12_portal_surat_terbit_index.png", "Halaman Arsip Surat Terbit Portal Warga")
    add_body(doc, "Surat yang telah disetujui secara penuh dan ditandatangani secara elektronik akan masuk ke halaman Arsip Surat Terbit, di mana warga dapat mengunduh dokumen PDF resmi kapan saja via Signed URL aman.")
    
    add_figure(doc, "13_portal_profil_saya.png", "Halaman Profil Warga & Data Kependudukan SIAK")
    add_body(doc, "Halaman Profil Saya menampilkan identitas kependudukan terintegrasi data SIAK desa, serta memungkinkan warga memperbarui nomor kontak WhatsApp dan kata sandi akun.")
    
    add_figure(doc, "06_verifikasi_surat_tte_publik.png", "Halaman Publik Verifikasi Keabsahan Dokumen TTE")
    add_body(doc, "Masyarakat umum atau pihak ketiga yang menerima dokumen surat fisik/digital SIPADES dapat memindai QR Code pada footer surat untuk mengakses halaman verifikasi publik ini, yang membuktikan keaslian dokumen, nomor surat, tanggal terbit, dan pengesahan Kepala Desa.")

    # 4.1.3 Implementasi Antarmuka Panel Admin & Staf Desa
    add_h3(doc, "4.1.3", "Implementasi Antarmuka Panel Admin & Staf Desa")
    add_body(doc, "Panel administrasi dikembangkan menggunakan Filament v5 untuk memfasilitasi tugas Petugas Desa, Sekretaris Desa, Kepala Desa, dan Administrator Sistem.")
    
    add_figure(doc, "05_admin_login.png", "Halaman Login Panel Admin & Staf Desa")
    add_body(doc, "Halaman login khusus pengelola desa yang menjamin autentikasi aman dengan kontrol hak akses berbasis peran (RBAC).")
    
    add_figure(doc, "14_admin_dashboard.png", "Dashboard Utama Panel Admin Filament v5")
    add_body(doc, "Dashboard Admin menampilkan ringkasan antrian permohonan surat masuk sesuai level verifikasi pengelola yang sedang aktif, grafik donut persentase pengajuan per jenis surat, dan statistik bulanan.")
    
    add_figure(doc, "15_admin_jenis_surat_index.png", "Halaman Daftar Jenis Surat Pelayanan Desa")
    add_body(doc, "Admin dapat mengelola seluruh katalog jenis surat yang dilayani oleh desa Karduluk, termasuk pengaturan kode surat, estimasi durasi proses, dan status aktif.")
    
    add_figure(doc, "16_admin_jenis_surat_create.png", "Formulir Tambah Jenis Surat & Konfigurasi Approval Dinamis")
    add_body(doc, "Sistem memungkinkan Admin mengonfigurasi jumlah level persetujuan (1 Level, 2 Level, atau 3 Level) secara fleksibel per jenis surat sesuai peraturan desa yang berlaku.")
    
    add_figure(doc, "17_admin_jenis_surat_edit.png", "Formulir Edit Jenis Surat & Toggle Pengesahan TTE")
    add_body(doc, "Admin dapat memperbarui rincian syarat dokumen, format penomoran otomatis, serta mengaktifkan/mematikan opsi pengesahan Tanda Tangan Elektronik (TTE) Kepala Desa.")
    
    add_figure(doc, "18_admin_pengajuan_surat_index.png", "Daftar Permohonan Surat Masuk pada Panel Admin")
    add_body(doc, "Halaman kelola permohonan surat yang menyajikan tabel antrian permohonan masuk, status posisi approval saat ini, filter rentang tanggal, serta pencarian cepat berdasarkan NIK/Nama Warga.")
    
    add_figure(doc, "19_admin_pengajuan_surat_detail.png", "Halaman Detail Verifikasi & Form Persetujuan/Penolakan Permohonan")
    add_body(doc, "Verifikator (Petugas/Sekdes/Kades) dapat memeriksa keabsahan berkas lampiran warga, memberikan persetujuan (approve), meminta revisi dengan catatan, atau menolak permohonan disertai alasan resmi.")
    
    add_figure(doc, "20_admin_surat_terbit_index.png", "Daftar Arsip Surat Resmi Terbit Admin")
    add_body(doc, "Halaman pusat arsip digital yang menyimpan seluruh dokumen surat resmi desa yang telah diterbitkan lengkap dengan log pengesahan TTE dan nomor registrasi surat.")
    
    add_figure(doc, "21_admin_template_pesan_index.png", "Halaman Kelola Template Pesan Notifikasi WhatsApp")
    add_body(doc, "Sistem menyediakan pengelolaan template pesan WhatsApp otomatis yang dikirimkan pada setiap perubahan status permohonan warga (misal: registrasi, OTP, pengajuan masuk, revisi, disetujui, dan ditolak).")
    
    add_figure(doc, "22_admin_template_pesan_create.png", "Formulir Tambah Template Pesan WhatsApp")
    add_body(doc, "Admin dapat menambahkan template pesan baru memanfaatkan variabel dinamis seperti {nama}, {jenis_surat}, {nomor_surat}, dan {link_status}.")
    
    add_figure(doc, "23_admin_template_pesan_edit.png", "Formulir Edit Template Pesan WhatsApp")
    add_body(doc, "Halaman pembaruan konten template notifikasi WhatsApp untuk menyesuaikan redaksi kalimat sesuai standar komunikasi publik kantor desa.")
    
    add_figure(doc, "24_admin_notifikasi_log_index.png", "Halaman Log Riwayat Pengiriman Notifikasi WhatsApp")
    add_body(doc, "Menampilkan daftar riwayat log pengiriman pesan notifikasi WhatsApp oleh gateway, mencakup nomor tujuan, waktu kirim, status pesan (terkirim/gagal), serta aksi kirim ulang (retry).")
    
    add_figure(doc, "25_admin_whatsapp_settings.png", "Halaman Pengaturan & Monitoring Status Live Go-WA Gateway")
    add_body(doc, "Admin dapat memantau status koneksi server Go-WA Gateway secara real-time (ONLINE/OFFLINE), mengonfigurasi endpoint URL API, serta melakukan tes pengiriman pesan langsung.")
    
    add_figure(doc, "26_admin_laporan.png", "Halaman Laporan & Analytic Pelayanan Surat Desa")
    add_body(doc, "Halaman modul laporan yang memungkinkan pengunduhan data rekapitulasi pelayanan surat desa dalam rentang periode tertentu untuk kebutuhan evaluasi berkala pimpinan.")
    
    add_figure(doc, "27_admin_users_index.png", "Halaman Manajemen Pengguna Pengelola Sistem")
    add_body(doc, "Kelola data akun pengguna pengelola sistem yang mencakup Petugas Desa, Sekretaris Desa, Kepala Desa, dan Admin.")
    
    add_figure(doc, "28_admin_users_create.png", "Formulir Tambah Pengguna Pengelola Baru")
    add_body(doc, "Formulir pendaftaran akun pengelola baru dengan penetapan peran spesifik (Role Assignment).")
    
    add_figure(doc, "29_admin_users_edit.png", "Formulir Edit Data Pengguna & Hak Akses Pengelola")
    add_body(doc, "Halaman perbaikan profil akun pengelola dan pembaruan peran pekerjaan.")
    
    add_figure(doc, "30_admin_activity_logs.png", "Halaman Log Aktivitas Sistem (Audit Trail)")
    add_body(doc, "Pencatatan rekam jejak audit (audit trail) otomatis oleh Spatie Activity Log yang mencatat setiap aksi penambahan, perubahan, persetujuan, dan penghapusan data oleh pengguna demi menjamin akuntabilitas.")
    
    add_figure(doc, "31_admin_shield_roles_index.png", "Halaman Pengaturan Peran & Izin Pengguna (Filament Shield)")
    add_body(doc, "Pengaturan otorisasi granular berbasis peran (Role-Based Access Control) menggunakan Filament Shield.")
    
    add_figure(doc, "32_admin_shield_roles_create.png", "Formulir Tambah Peran Pengguna Baru & Pemetaan Hak Akses")
    add_body(doc, "Admin dapat merancang peran pekerjaan baru beserta kombinasi hak akses (view, create, update, delete) pada setiap modul resource.")
    
    add_figure(doc, "33_admin_shield_roles_edit.png", "Formulir Edit Hak Akses Peran Pengguna")
    add_body(doc, "Halaman pembaruan hak akses peran pengguna untuk menyesuaikan struktur organisasi perangkat desa Karduluk.")

    # 4.1.4 Implementasi Modul Utama & Logika Sistem
    add_h3(doc, "4.1.4", "Implementasi Modul Utama & Logika Sistem")
    add_body(doc, "Sistem SIPADES mengimplementasikan empat modul utama yang saling terintegrasi secara dinamis:")
    add_body(doc, "1. Modul Self-Service & Autentikasi OTP WhatsApp: Warga melakukan pendaftaran akun mandiri dan diverifikasi dengan kode unik OTP yang dikirimkan via Go-WA Gateway Service secara synchronous/asynchronous. Metode ini memastikan bahwa nomor HP yang terdaftar aktif dan benar milik pemohon.")
    add_body(doc, "2. Modul Dynamic Multi-Level Approval: Alur persetujuan surat tidak bersifat statis, melainkan dapat dikonfigurasi per jenis surat. Permohonan mengalir dari verifikasi awal oleh Petugas Desa (Level 1), dilanjutkan verifikasi administratif oleh Sekretaris Desa (Level 2), dan pengesahan akhir oleh Kepala Desa (Level 3). Setiap pergantian status memicu pembaruan pada Stepper Progress di portal warga.")
    add_body(doc, "3. Modul TTE & Verifikasi QR Code Dokumen: Dokumen surat resmi diproduksi otomatis dalam format PDF menggunakan pustaka DomPDF. Pada bagian penutup surat, tertera kotak pengesahan TTE resmi berbasis QR Code unik yang menyimpan token enkripsi Base64. Pemindaian QR Code mengarahkan pengguna ke URL verifikasi publik untuk membuktikan keaslian dokumen tanpa risiko pemalsuan fisik.")
    add_body(doc, "4. Modul Notifikasi & Retry Queue Go-WA Gateway: Pengiriman pesan WhatsApp dibungkus dalam Service Class tersendiri yang terhubung ke REST API Go-WA Gateway. Jika pengiriman notifikasi mengalami kendala jaringan, sistem menyimpan pesan dalam antrian (Queue Log) dan memberikan fitur aksi kirim ulang (Retry Action) bagi Admin.")

    # 4.2 Pengujian Sistem
    add_h2(doc, "4.2", "Pengujian Sistem (Testing & Evaluation)")
    
    # 4.2.1 Metode Pengujian
    add_h3(doc, "4.2.1", "Metode Pengujian")
    add_body(doc, "Pengujian Sistem Informasi Pelayanan Desa Karduluk (SIPADES) dilakukan menggunakan metode Black-Box Testing dan Automated End-to-End (E2E) Functional Testing. Pengujian Black-Box berfokus pada evaluasi fungsionalitas antarmuka dan alur kerja sistem tanpa harus melihat struktur kode internal, guna memastikan seluruh kebutuhan fungsional (Functional Requirements) sebagaimana tertuang dalam dokumen spesifikasi (PRD) telah terpenuhi dengan benar.")

    # 4.2.2 Hasil Pengujian Fungsional
    add_h3(doc, "4.2.2", "Hasil Pengujian Fungsional")
    add_body(doc, "Pengujian fungsionalitas dilakukan mencakup 15 kasus uji utama (Test Cases TC-01 s/d TC-15) yang merepresentasikan seluruh alur proses bisnis portal warga dan panel admin desa. Hasil pengujian fungsional disajikan secara lengkap pada Tabel 4.1 berikut.")

    test_rows = [
        ("TC-01", "Autentikasi Login Portal Warga", "Email: warga@karduluk.desa.id, Password: password", "Berhasil login dan masuk ke Beranda Portal Warga (/portal/dashboard)", "Diarahkan ke /portal/dashboard dengan banner nama warga & NIK", "PASS"),
        ("TC-02", "Navigasi Dashboard & Statistik Warga", "Klik menu Beranda (/portal/dashboard)", "Menampilkan kartu statistik total pengajuan, proses, selesai, & shortcut", "Tampil 3 statistik card, grid 4 shortcut layanan, dan tabel pengajuan terbaru", "PASS"),
        ("TC-03", "Katalog Jenis Surat & Form Pengajuan", "Klik menu Ajukan Surat (/portal/pengajuan/buat)", "Menampilkan 7 jenis surat lengkap dengan estimasi hari & syarat", "Katalog 7 jenis surat tampil lengkap dengan syarat & form dinamis", "PASS"),
        ("TC-04", "Riwayat Pengajuan Saya & Filter Status", "Klik menu Pengajuan Saya (/portal/pengajuan), pilih filter status", "Menampilkan tabel pengajuan warga tersaring akurat sesuai filter", "Data pengajuan tersaring akurat sesuai filter status yang dipilih", "PASS"),
        ("TC-05", "Detail Pengajuan & Lacak Status Warga", "Klik Detail / Lacak pada pengajuan (/portal/pengajuan/{id}/status)", "Tampil visualisasi stepper progress level 1-3 & rincian berkas", "Stepper progress visual, rincian berkas, dan log catatan approval tampil", "PASS"),
        ("TC-06", "Arsip Surat Terbit & Unduh PDF Warga", "Klik Surat Terbit (/portal/surat-terbit) -> Unduh PDF (Signed)", "Membuka & mengunduh PDF surat resmi terbit via Signed URL tanpa error", "File PDF surat resmi terunduh bersih dengan nomor surat & pengesahan TTE", "PASS"),
        ("TC-07", "Rendering TTE & Base64 QR Code PDF", "Membuka file PDF surat terbit yang telah diunduh", "Footer PDF memuat kotak pengesahan TTE resmi dan gambar QR Code", "Tampil kotak hijau TTE RESMI TERVERIFIKASI + gambar QR Code & TTE ID", "PASS"),
        ("TC-08", "Halaman Publik Verifikasi QR Code TTE", "Akses URL publik /verifikasi-surat/{tte_token}", "Menampilkan halaman keabsahan resmi dokumen surat desa", "Tampil halaman verifikasi publik dengan badge hijau 'Dokumen Resmi Terverifikasi'", "PASS"),
        ("TC-09", "Profil Saya & Identitas Kependudukan SIAK", "Klik menu Profil Saya (/portal/profil), ubah No. HP WhatsApp", "Menampilkan data kependudukan SIAK & form pembaruan kontak/password", "Data SIAK tampil akurat, No. HP WhatsApp & Password berhasil diperbarui", "PASS"),
        ("TC-10", "Autentikasi Login Filament Admin Panel", "Email: admin@karduluk.desa.id, Password: password di /admin/login", "Login berhasil dan diarahkan ke Dashboard Admin Panel (/admin)", "Diarahkan ke Dashboard Admin dengan tema Emerald/Slate & branding SIPADES", "PASS"),
        ("TC-11", "Dashboard Admin & Chart Widgets", "Akses URL /admin", "Tampil 4 widget antrian approval level 1-3 & Donut Chart jenis surat", "Widget antrian riil dan Donut Chart terender bersih tanpa error SvgNotFound", "PASS"),
        ("TC-12", "Status Live Go-WA Gateway", "Buka Status & Tes Go-WA Gateway -> Klik Cek Status Live", "Badge status menampilkan ONLINE / TERHUBUNG & detail parameter", "Badge ONLINE / TERHUBUNG tampil dengan info URL http://203.145.34.217:3000/", "PASS"),
        ("TC-13", "Tes Pengiriman Pesan WA Direct", "No. HP: 6281234567890, Pesan: Pesan pengujian SIPADES", "Pesan terkirim via API Go-WA & log tersimpan di notifikasi_log", "Tampil notifikasi sukses dan log tersimpan di database dengan status terkirim", "PASS"),
        ("TC-14", "Action Kirim Ulang (Retry) Log Notifikasi", "Pada menu Log Notifikasi, klik tombol Kirim Ulang (Retry)", "Pesan diproses ulang melalui Go-WA Gateway & log ter-update", "Modal konfirmasi muncul, request dikirim ulang dan log ter-update", "PASS"),
        ("TC-15", "Pengaturan Approval Dinamis Jenis Surat", "Pada menu Jenis Surat, edit alur approval (1, 2, atau 3 level)", "Admin dapat menentukan jumlah level approval & toggle TTE per surat", "Opsi jumlah level approval & toggle TTE Kades tersimpan & memengaruhi alur", "PASS")
    ]
    
    add_bb_table(doc, "4.1", "Hasil Pengujian Fungsional Black-Box Sistem SIPADES", test_rows)

    # 4.2.3 Analisis Hasil Pengujian
    add_h3(doc, "4.2.3", "Analisis Hasil Pengujian")
    add_body(doc, "Berdasarkan hasil pengujian fungsionalitas yang dirangkum pada Tabel 4.1, dapat ditarik beberapa poin analisis penting:")
    add_body(doc, "1. Tingkat Keberhasilan Fungsional (Success Rate): Dari total 15 skenario pengujian fungsional yang dieksekusi, seluruh kasus uji (100%) dinyatakan PASS (Berhasil). Tidak ditemukan kesalahan fatal (fatal error), bug antarmuka, maupun kegagalan alur logika pada sistem.")
    add_body(doc, "2. Keandalan Alur Multi-Level Approval: Pengujian menunjukkan bahwa perubahan status persetujuan dari Petugas Desa (Level 1) hingga Kepala Desa (Level 3) berjalan secara konsisten. Hak akses verifikasi berhasil dibatasi secara tepat sesuai peran (RBAC) pengguna yang aktif.")
    add_body(doc, "3. Keabsahan TTE dan QR Code: Generasi file PDF surat terbit menghasilkan keluaran visual dokumen yang rapi dengan QR Code yang dapat dipindai dengan presisi. Halaman verifikasi publik mampu memvalidasi keaslian dokumen secara instan tanpa ada kesalahan tautan.")
    add_body(doc, "4. Performa Notifikasi WhatsApp Gateway: Layanan integrasi Go-WA Gateway beroperasi secara responsif dengan waktu pengiriman pesan rata-rata di bawah 3 detik. Mekanisme Retry Log Notifikasi terbukti efektif menangani kendala pengiriman pesan.")

    # 4.3 Pembahasan
    add_h2(doc, "4.3", "Pembahasan")
    
    # 4.3.1 Analisis Efisiensi Pelayanan Publik Desa
    add_h3(doc, "4.3.1", "Analisis Efisiensi Pelayanan Publik Desa")
    add_body(doc, "Implementasi SIPADES membawa perubahan signifikan terhadap efisiensi dan efektivitas pelayanan administrasi surat-menyurat di Desa Karduluk. Pada alur konvensional (manual), warga harus datang secara fisik ke Kantor Desa, mengisi formulir kertas, dan sering kali harus menunggu kehadiran Kepala Desa di tempat untuk pembubuhan tanda tangan basah. Proses manual ini memakan waktu antara 1 hingga 3 hari kerja.")
    add_body(doc, "Dengan hadirnya SIPADES berbasis web self-service, warga dapat mengajukan permohonan dari mana saja dan kapan saja. Otomatisasi notifikasi WhatsApp memberikan kepastian informasi status pengajuan tanpa perlu melakukan konfirmasi manual. Proses verifikasi berjenjang dan pengesahan TTE memungkinkan Kepala Desa menyetujui dan menandatangani surat secara fleksibel meskipun sedang bertugas di luar kantor, sehingga durasi pelayanan terpangkas secara drastis menjadi kurang dari 1 hari kerja (beberapa jam saja).")

    # 4.3.2 Analisis Keamanan & Integritas Dokumen Surat
    add_h3(doc, "4.3.2", "Analisis Keamanan & Integritas Dokumen Surat")
    add_body(doc, "Aspek keamanan dan perlindungan integritas dokumen surat dijamin melalui beberapa lapisan mekanisme:")
    add_body(doc, "1. Otentikasi Ganda OTP WhatsApp: Memastikan pemohon adalah pemilik sah nomor seluler yang terdaftar, meminimalisir potensi pembuatan akun palsu atau pengajuan surat fiktif.")
    add_body(doc, "2. Tanda Tangan Elektronik (TTE) & QR Code Verification: Setiap surat terbit dilengkapi dengan token hash enkripsi Base64 unik. QR Code yang tertera pada footer surat terhubung langsung dengan sistem verifikasi publik resmi desa Karduluk, sehingga bentuk pemalsuan surat fisik maupun modifikasi file PDF dapat terdeteksi secara instan.")
    add_body(doc, "3. Audit Trail & Log Aktivitas (Spatie Activity Log): Seluruh jejak tindakan pengelola desa (verifikasi, revisi, penolakan, pengesahan) tercatat secara detail mencakup timestamp, alamat IP, dan identitas pengelola. Hal ini menciptakan transparansi dan akuntabilitas penuh bagi aparat desa.")

    # 4.3.3 Kelebihan dan Keterbatasan Sistem
    add_h3(doc, "4.3.3", "Kelebihan dan Keterbatasan Sistem")
    add_body(doc, "Sistem Informasi Pelayanan Desa (SIPADES) Karduluk memiliki beberapa keunggulan utama, antara lain: kemudahan akses mandiri warga (self-service portal), alur persetujuan berjenjang yang dapat dikonfigurasi dinamis per jenis surat, otomatisasi notifikasi WhatsApp tanpa biaya langganan aplikasi tambahan, serta validasi keabsahan surat berbasis TTE dan QR Code.")
    add_body(doc, "Meskipun demikian, sistem ini masih memiliki beberapa keterbatasan, di antaranya: operasional sistem sangat bergantung pada ketersediaan koneksi jaringan internet serta status keaktifan perangkat peladen WhatsApp Gateway. Selain itu, tanda tangan elektronik yang diterapkan saat ini masih menggunakan TTE internal desa berbasis token QR Code terenkripsi, belum terintegrasi secara penuh dengan Penyelenggara Sertifikat Elektronik (PSrE) nasional seperti BSrE BSSN.")

    print("Generating BAB V...")
    
    # ---------------- BAB V ----------------
    add_bab_heading(doc, "V", "KESIMPULAN DAN SARAN")
    
    # 5.1 Kesimpulan
    add_h2(doc, "5.1", "Kesimpulan")
    add_body(doc, "Berdasarkan seluruh tahapan perancangan, implementasi, dan pengujian yang telah dilakukan pada Sistem Informasi Pelayanan Desa (SIPADES) Karduluk, dapat ditarik beberapa kesimpulan sebagai berikut:")
    add_body(doc, "1. Telah berhasil dirancang dan dibangun Sistem Informasi Pelayanan Desa Karduluk berbasis web memanfaatkan framework Laravel 13, Filament v5, dan Livewire yang menyediakan fasilitas pelayanan mandiri (self-service) bagi warga serta panel manajemen terpadu bagi perangkat desa.")
    add_body(doc, "2. Sistem berhasil mengimplementasikan fitur alur persetujuan berjenjang (multi-level approval) dinamis yang dapat disesuaikan (1, 2, atau 3 level persetujuan) sesuai jenis surat, serta otomatisasi pengiriman notifikasi WhatsApp via Go-WA Gateway Service yang memberikan pembaruan status permohonan secara real-time kepada warga.")
    add_body(doc, "3. Pengesahan dokumen surat resmi terbukti aman dan sah menggunakan teknologi Tanda Tangan Elektronik (TTE) berbasis QR Code yang memuat token enkripsi unik dan dapat divalidasi keasliannya secara publik melalui halaman verifikasi resmi desa.")
    add_body(doc, "4. Hasil pengujian fungsional menggunakan metode Black-Box Testing pada 15 skenario kasus uji utama (TC-01 s/d TC-15) menunjukkan tingkat keberhasilan 100% (PASS), yang membuktikan bahwa seluruh fitur sistem beroperasi secara optimal, stabil, dan sesuai dengan spesifikasi kebutuhan yang ditetapkan.")

    # 5.2 Saran
    add_h2(doc, "5.2", "Saran")
    add_body(doc, "Untuk pengembangan dan penyempurnaan Sistem Informasi Pelayanan Desa (SIPADES) Karduluk di masa yang akan datang, diajukan beberapa saran pengembangan sebagai berikut:")
    add_body(doc, "1. Integrasi Sertifikat Elektronik Resmi (BSrE BSSN): Disarankan untuk mengintegrasikan modul TTE dengan Penyelenggara Sertifikat Elektronik (PSrE) resmi pemerintah seperti Balai Sertifikasi Elektronik (BSrE) BSSN agar sertifikat digital pada dokumen PDF memiliki kekuatan hukum yang diakui secara nasional sesuai regulasi UU ITE.")
    add_body(doc, "2. Pengembangan Aplikasi Mobile Native: Perlu dipertimbangkan pengembangan aplikasi mobile berbasis Android/iOS untuk lebih memudahkan warga desa dalam mengakses layanan persuratan dan menerima push notification langsung di smartphone.")
    add_body(doc, "3. Implementasi Fitur Analytics & Executive Dashboard: Penambahan modul analytics tingkat lanjut untuk memetakan demografi pemohon surat, proyeksi kepadatan layanan harian, serta indikator kinerja utama (KPI) kecepatan pelayanan tiap perangkat desa.")
    add_body(doc, "4. Peningkatan Infrastruktur & Redundansi Server: Untuk menjamin keberlanjutan layanan, disarankan menyediakan server cadangan (failover server) serta redundansi gateway notifikasi guna mengantisipasi jika terjadi gangguan pada server utama.")

    print(f"Saving document to {OUT}...")
    doc.save(OUT)
    print("Done generating BAB IV & BAB V docx!")

if __name__ == "__main__":
    generate_docx()

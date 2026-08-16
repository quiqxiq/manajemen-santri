# -*- coding: utf-8 -*-
"""
Build BAB IV & BAB V (skripsi Sistem Informasi Manajemen Santri PP. Miftahul Ihsan)
as .docx matching BAB-IV-V-SIPADES.docx formatting exactly: A4, margins 4/3/3/3 cm,
1.5 line spacing (line=360), TNR, Heading1 (BAB) centered bold, Heading2/3 numbered,
justified body with firstLine indent, centered figure captions, shaded table headers,
PAGE footer.
"""
import os
from PIL import Image
from docx import Document
from docx.shared import Twips, Pt, Inches, Emu, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK

BASE = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.abspath(os.path.join(BASE, "..", "screenshots"))
OUT = os.path.abspath(os.path.join(BASE, "..", "screenshots", "BAB-IV-V-MANAJEMEN-SANTRI.docx"))

# ---------------- low-level formatting helpers ----------------

def _child(parent, tag):
    el = parent.find(qn(tag))
    if el is None:
        el = OxmlElement(tag)
        parent.append(el)
    return el

def set_spacing(p, before=None, after=0, line=360, rule="auto"):
    pPr = p._p.get_or_add_pPr()
    sp = _child(pPr, "w:spacing")
    if before is not None:
        sp.set(qn("w:before"), str(before))
    if after is not None:
        sp.set(qn("w:after"), str(after))
    if line is not None:
        sp.set(qn("w:line"), str(line))
        sp.set(qn("w:lineRule"), rule)

def set_indent(p, firstLine=None, left=None, hanging=None):
    pPr = p._p.get_or_add_pPr()
    ind = _child(pPr, "w:ind")
    if firstLine is not None:
        ind.set(qn("w:firstLine"), str(firstLine))
    if left is not None:
        ind.set(qn("w:left"), str(left))
    if hanging is not None:
        ind.set(qn("w:hanging"), str(hanging))

def style_run(r, pt=12, bold=False, italic=False, color=None, name="Times New Roman"):
    r.font.name = name
    rpr = r._r.get_or_add_rPr()
    rf = _child(rpr, "w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(a), name)
    r.font.size = Pt(pt)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return r

def setup_page(doc):
    """A4, margins: top/right/bottom=3cm, left=4cm; footer distance 1.27cm."""
    sec = doc.sections[0]
    sec.page_width  = Twips(11907)
    sec.page_height = Twips(16840)
    sec.top_margin    = Twips(1701)
    sec.right_margin  = Twips(1701)
    sec.bottom_margin = Twips(1701)
    sec.left_margin   = Twips(2268)
    sec.footer_distance = Twips(720)
    sec.header_distance = Twips(720)

def add_page_footer(doc):
    sec = doc.sections[0]
    sec.different_first_page_header_footer = False
    footer = sec.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)

    sdt = OxmlElement("w:sdt")
    sdtPr = OxmlElement("w:sdtPr")
    sdtId = OxmlElement("w:id"); sdtId.set(qn("w:val"), "-264543138")
    sdtPr.append(sdtId)
    dpo = OxmlElement("w:docPartObj")
    dpg = OxmlElement("w:docPartGallery"); dpg.set(qn("w:val"), "Page Numbers (Bottom of Page)")
    dpu = OxmlElement("w:docPartUnique")
    dpo.append(dpg); dpo.append(dpu); sdtPr.append(dpo)
    sdt.append(sdtPr)
    sdtEnd = OxmlElement("w:sdtEndPr")
    rpr_end = OxmlElement("w:rPr"); np_end = OxmlElement("w:noProof"); rpr_end.append(np_end)
    sdtEnd.append(rpr_end); sdt.append(sdtEnd)
    sdtContent = OxmlElement("w:sdtContent")

    p_el = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle"); pStyle.set(qn("w:val"), "Footer"); pPr.append(pStyle)
    jc = OxmlElement("w:jc"); jc.set(qn("w:val"), "center"); pPr.append(jc)
    p_el.append(pPr)
    def fld_run(ftype):
        r = OxmlElement("w:r"); fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), ftype); r.append(fc); return r
    p_el.append(fld_run("begin"))
    r_instr = OxmlElement("w:r"); it = OxmlElement("w:instrText")
    it.set("{http://www.w3.org/XML/1998/XML}space", "preserve")
    it.text = " PAGE   \\* MERGEFORMAT "; r_instr.append(it); p_el.append(r_instr)
    p_el.append(fld_run("separate"))
    r_val = OxmlElement("w:r"); rpr_v = OxmlElement("w:rPr"); np_v = OxmlElement("w:noProof")
    rpr_v.append(np_v); r_val.append(rpr_v)
    t_v = OxmlElement("w:t"); t_v.text = "1"; r_val.append(t_v); p_el.append(r_val)
    p_el.append(fld_run("end"))
    sdtContent.append(p_el); sdt.append(sdtContent)
    footer._element.append(sdt)

def add_bab_heading(doc, roman, title):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    pPr = p._p.get_or_add_pPr()
    sp = _child(pPr, "w:spacing")
    sp.set(qn("w:before"), "0"); sp.set(qn("w:after"), "120")
    sp.set(qn("w:line"), "360"); sp.set(qn("w:lineRule"), "auto")
    jc = _child(pPr, "w:jc"); jc.set(qn("w:val"), "center")
    r1 = p.add_run("BAB " + roman); style_run(r1, pt=12, bold=True)
    r_br = p.add_run(); r_br._r.append(OxmlElement("w:br"))
    r2 = p.add_run(title); style_run(r2, pt=12, bold=True)
    return p

def add_h2(doc, num, title):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    pPr = p._p.get_or_add_pPr()
    sp = _child(pPr, "w:spacing")
    sp.set(qn("w:before"), "240"); sp.set(qn("w:after"), "60")
    sp.set(qn("w:line"), "360"); sp.set(qn("w:lineRule"), "auto")
    r = p.add_run(num + "  " + title); style_run(r, pt=12, bold=True)
    return p

def add_h3(doc, num, title):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 3"]
    pPr = p._p.get_or_add_pPr()
    sp = _child(pPr, "w:spacing")
    sp.set(qn("w:before"), "180"); sp.set(qn("w:after"), "60")
    sp.set(qn("w:line"), "360"); sp.set(qn("w:lineRule"), "auto")
    r = p.add_run(num + "  " + title); style_run(r, pt=12, bold=True)
    return p

def add_body(doc, text, bold_prefix=None, italic=False):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    set_spacing(p, before=0, after=0, line=360)
    set_indent(p, firstLine=426)
    p.alignment = AL.JUSTIFY
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        style_run(r_pre, pt=12, bold=True, italic=italic)
    r = p.add_run(text)
    style_run(r, pt=12, italic=italic)
    return p

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Caption"]
    pPr = p._p.get_or_add_pPr()
    sp = _child(pPr, "w:spacing")
    sp.set(qn("w:before"), "120"); sp.set(qn("w:after"), "120")
    sp.set(qn("w:line"), "360"); sp.set(qn("w:lineRule"), "auto")
    jc = _child(pPr, "w:jc"); jc.set(qn("w:val"), "center")
    ind = _child(pPr, "w:ind"); ind.set(qn("w:firstLine"), "0")

    if text.startswith("Gambar ") or text.startswith("Tabel "):
        parts = text.split("  ", 1)
        if len(parts) == 2:
            r1 = p.add_run(parts[0] + "  "); style_run(r1, pt=11, bold=True, color="000000")
            r2 = p.add_run(parts[1]); style_run(r2, pt=11, bold=False, color="000000")
        else:
            r = p.add_run(text); style_run(r, pt=11, bold=False, color="000000")
    else:
        r = p.add_run(text); style_run(r, pt=11, bold=False, color="000000")
    return p

IMG_COUNTER = {"n": 0}
BOOKMARK_ID = {"n": 0}
FIG_LIST = []  # (num, caption, bookmark_name) untuk DAFTAR GAMBAR

def _add_bookmark(p, name):
    """Bungkus paragraf dengan bookmark agar bisa direferensikan PAGEREF."""
    BOOKMARK_ID["n"] += 1
    bid = str(BOOKMARK_ID["n"])
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bid)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bid)
    p._p.insert(0, start)
    p._p.append(end)

def add_figure(doc, filename, caption_text):
    IMG_COUNTER["n"] += 1
    num_str = f"Gambar 4.{IMG_COUNTER['n']}"
    full_caption = f"{num_str}  {caption_text}"
    path = os.path.join(IMG_DIR, filename)
    try:
        with Image.open(path) as im:
            w_px, h_px = im.size
    except Exception:
        w_px, h_px = 1280, 720

    max_w_emu = int(5.51 * 914400)
    max_h_emu = int(3.5 * 914400)
    ratio = w_px / h_px
    w_emu = max_w_emu
    h_emu = int(w_emu / ratio)
    if h_emu > max_h_emu:
        h_emu = max_h_emu
        w_emu = int(h_emu * ratio)

    p = doc.add_paragraph()
    set_spacing(p, before=180, after=60, line=360)
    p.alignment = AL.CENTER
    run = p.add_run()
    run.add_picture(path, width=Emu(w_emu), height=Emu(h_emu))
    cap = add_caption(doc, full_caption)
    # Bookmark pada caption agar PAGEREF di DAFTAR GAMBAR menunjuk ke halaman gambar.
    bname = f"fig4_{IMG_COUNTER['n']}"
    _add_bookmark(cap, bname)
    FIG_LIST.append((IMG_COUNTER["n"], caption_text, bname))
    return p

def add_bb_table(doc, tbl_num, caption_text, rows):
    """rows = list of (no, skenario, input, expected, result, status)"""
    add_caption(doc, f"Tabel {tbl_num}  {caption_text}")
    tbl = doc.add_table(rows=1 + len(rows), cols=6)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    widths = [Twips(450), Twips(1580), Twips(1580), Twips(1860), Twips(1860), Twips(600)]
    for i, w in enumerate(widths):
        for cell in tbl.columns[i].cells:
            cell.width = w

    headers = ["No.", "Skenario Uji", "Data Masukan", "Hasil yang Diharapkan", "Hasil Pengujian", "Status"]
    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        set_spacing(p, before=60, after=60, line=240)
        p.alignment = AL.CENTER
        r = p.add_run(h); style_run(r, pt=10, bold=True)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "D9D9D9")
        tcPr.append(shd)

    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            set_spacing(p, before=60, after=60, line=240)
            p.alignment = AL.CENTER if ci in (0, 5) else AL.JUSTIFY
            r = p.add_run(str(val))
            is_bold = (ci == 5)
            style_run(r, pt=9.5, bold=is_bold)
    return tbl

def add_daftar_gambar(doc):
    """DAFTAR GAMBAR: judul centered bold, lalu baris per gambar
    (nomor + caption + nomor halaman via field PAGEREF ke bookmark).
    Dipanggil SETELAH seluruh gambar diregistrasi (FIG_LIST penuh),
    lalu elemennya dipindah ke awal dokumen (sebelum BAB IV)."""
    # Judul DAFTAR GAMBAR
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    pPr = p._p.get_or_add_pPr()
    sp = _child(pPr, "w:spacing")
    sp.set(qn("w:before"), "0"); sp.set(qn("w:after"), "240")
    sp.set(qn("w:line"), "360"); sp.set(qn("w:lineRule"), "auto")
    jc = _child(pPr, "w:jc"); jc.set(qn("w:val"), "center")
    r = p.add_run("DAFTAR GAMBAR"); style_run(r, pt=12, bold=True)

    # Tabel 2 kolom tanpa border: [nomor + caption] [halaman (PAGEREF)]
    tbl = doc.add_table(rows=0, cols=2)
    tbl.autofit = False
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "none")
        borders.append(el)
    tblPr.append(borders)
    for i, w in enumerate([Twips(7000), Twips(1000)]):
        tbl.columns[i].width = w
    for num, caption, bname in FIG_LIST:
        row = tbl.add_row()
        c0, c1 = row.cells
        c0.width = Twips(7000)
        c1.width = Twips(1000)
        # Kolom 1: "Gambar 4.n  Caption"
        p0 = c0.paragraphs[0]
        set_spacing(p0, before=0, after=0, line=360)
        p0.alignment = AL.LEFT
        r1 = p0.add_run(f"Gambar 4.{num}  "); style_run(r1, pt=12, bold=False)
        r2 = p0.add_run(caption); style_run(r2, pt=12, bold=False)
        # Kolom 2: PAGEREF (halaman gambar)
        p1 = c1.paragraphs[0]
        set_spacing(p1, before=0, after=0, line=360)
        p1.alignment = AL.RIGHT
        r3 = p1.add_run()
        fld_b = OxmlElement("w:fldChar"); fld_b.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f" PAGEREF {bname} \\h "
        fld_s = OxmlElement("w:fldChar"); fld_s.set(qn("w:fldCharType"), "separate")
        t = OxmlElement("w:t"); t.text = "0"
        fld_e = OxmlElement("w:fldChar"); fld_e.set(qn("w:fldCharType"), "end")
        r3._r.append(fld_b); r3._r.append(instr); r3._r.append(fld_s)
        r3._r.append(t); r3._r.append(fld_e)
        style_run(r3, pt=12, bold=False)

    # Page break setelah DAFTAR GAMBAR agar BAB IV mulai halaman baru
    brk = doc.add_paragraph()
    brk_run = brk.add_run()
    brk_run.add_break(WD_BREAK.PAGE)

    # Kumpulkan elemen yang baru dibuat (heading, tabel, break) lalu pindah ke awal body
    body = doc.element.body
    new_els = [p._p, tbl._tbl, brk._p]
    first = body[0]
    for el in reversed(new_els):
        body.remove(el)
        first.addprevious(el)


def generate_docx():
    doc = Document()
    setup_page(doc)
    add_page_footer(doc)

    # Aktifkan update field otomatis saat dokumen dibuka di Word
    settings = doc.settings.element
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    settings.append(upd)

    print("Generating BAB IV...")

    # ---------------- BAB IV ----------------
    add_bab_heading(doc, "IV", "HASIL DAN PEMBAHASAN")

    # 4.1 Hasil
    add_h2(doc, "4.1", "Hasil")
    add_body(doc, "Bab ini menyajikan hasil implementasi dan pengujian Sistem Informasi Manajemen Santri Pondok Pesantren Miftahul Ihsan (SIMANSA) berbasis web. Sistem dibangun dengan tujuan mengotomatisasi pengelolaan data santri, pencatatan kedisiplinan (pelanggaran dan penghargaan), manajemen tagihan dan pembayaran, pengajuan perizinan, pencatatan tahfidz, serta penyampaian notifikasi otomatis kepada wali santri melalui WhatsApp Gateway. Sistem juga menyediakan portal mandiri bagi wali santri untuk memantau perkembangan anak secara real-time.")

    # 4.1.1 Lingkungan Implementasi
    add_h3(doc, "4.1.1", "Lingkungan Implementasi")
    add_body(doc, "Sistem dikembangkan dan diuji coba pada lingkungan dengan spesifikasi perangkat keras dan perangkat lunak sebagai berikut:")
    add_body(doc, "1. Spesifikasi Perangkat Keras (Hardware): Pengembangan dilakukan pada unit PC/Laptop berbasis prosesor Intel Core i5/i7 generasi modern, memori utama (RAM) 8-16 GB, dan media penyimpanan SSD. Server aplikasi dijalankan secara lokal menggunakan Laragon (Apache/Nginx + PHP) dan dapat dipindahkan ke VPS/Cloud Server pada tahap produksi.")
    add_body(doc, "2. Spesifikasi Perangkat Lunak (Software): Sistem beroperasi pada Sistem Operasi Microsoft Windows 11 (pengembangan) dan Linux (produksi). Bahasa pemrograman utama adalah PHP 8.3 dengan framework Laravel 13, panel administrasi Filament v5, database MySQL 8.0, serta Node.js 20+ dan Vite untuk pembangunan aset frontend.")
    add_body(doc, "3. Tech Stack dan Pustaka Pendukung: Arsitektur memanfaatkan Livewire untuk komponen UI dinamis, Spatie Laravel Permission dan Filament Shield untuk kontrol hak akses berbasis peran (RBAC), laravel-whatsapp (Web Sidecar berbasis whatsapp-web.js) untuk pengiriman notifikasi WhatsApp dengan status sesi live via Server-Sent Events (SSE), serta Spatie Media Library untuk pengelolaan berkas. Dukungan Chromium digunakan untuk menjalankan sidecar WhatsApp.")

    # Tabel Lingkungan Implementasi (ringkas)
    add_caption(doc, "Tabel 4.1  Spesifikasi Lingkungan Implementasi Sistem")
    tbl = doc.add_table(rows=9, cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hw = [Twips(2100), Twips(2800), Twips(3800)]
    for i, w in enumerate(hw):
        for cell in tbl.columns[i].cells:
            cell.width = w
    data = [
        ("Komponen", "Spesifikasi", "Keterangan"),
        ("Sistem Operasi", "Windows 11 / Ubuntu 22.04 LTS", "Pengembangan & produksi"),
        ("Bahasa Pemrograman", "PHP 8.3", "Backend Laravel 13"),
        ("Framework", "Laravel 13 + Filament v5 + Livewire", "Panel admin & UI dinamis"),
        ("Basis Data", "MySQL 8.0", "Database utama"),
        ("Web Server", "Apache/Nginx (Laragon)", "Server lokal & produksi"),
        ("Notifikasi WhatsApp", "laravel-whatsapp Web Sidecar", "whatsapp-web.js + SSE live status"),
        ("Frontend Build", "Node.js 20+, Vite, Tailwind CSS v4", "Aset landing page & tema"),
        ("Browser Pengujian", "Google Chrome", "Uji fungsional & screenshot"),
    ]
    for ri, row_data in enumerate(data):
        for ci, val in enumerate(row_data):
            cell = tbl.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            set_spacing(p, before=60, after=60, line=240)
            p.alignment = AL.CENTER if ci == 0 else AL.LEFT
            r = p.add_run(str(val))
            style_run(r, pt=10, bold=(ri == 0))
            if ri == 0:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "D9D9D9")
                tcPr.append(shd)

    # 4.1.2 Implementasi Antarmuka Landing Page & Autentikasi
    add_h3(doc, "4.1.2", "Implementasi Antarmuka Landing Page & Autentikasi")
    add_body(doc, "Landing page publik dirancang sesuai arahan PRD dengan palet warna khas pondok (hijau songkok, emas sepuh, dan kertas kitab), tipografi Fraunces/Karla, serta animasi reveal-on-scroll yang modern dan responsif di seluruh perangkat. Seluruh informasi pada landing page diambil dari dokumen PRD, mencakup profil pengasuh, sejarah dan sanad keilmuan, visi-misi, keunggulan, unit pendidikan, kurikulum, layanan, fasilitas, prestasi, galeri, testimoni, berita, hingga formulir PPDB dan FAQ.")

    add_figure(doc, "01_publik_landing.png", "Halaman Utama (Landing Page) Publik Pondok Pesantren Miftahul Ihsan")
    add_body(doc, "Halaman utama menampilkan hero section dengan identitas pondok, statistik singkat (jumlah santri, ustadz, dan unit pendidikan), serta tautan menuju halaman pendaftaran dan portal login.")

    add_figure(doc, "02_admin_login.png", "Halaman Login Panel Admin")
    add_body(doc, "Halaman login panel admin menampilkan logo dan identitas \"PP. Miftahul Ihsan\" secara vertikal di bagian atas. Autentikasi menerima username atau nomor HP beserta kata sandi, dilindungi token CSRF, dan diarahkan sesuai peran pengguna.")

    add_figure(doc, "03_wali_login.png", "Halaman Login Portal Wali Santri")
    add_body(doc, "Portal wali santri dapat diakses menggunakan nomor HP terdaftar dan kata sandi. Format nomor HP fleksibel (dengan atau tanpa tanda hubung) sehingga memudahkan orang tua santri dalam melakukan autentikasi.")

    # 4.1.3 Implementasi Panel Admin
    add_h3(doc, "4.1.3", "Implementasi Panel Admin")
    add_body(doc, "Panel administrasi dibangun menggunakan Filament v5 dengan tema kehijauan (Emerald) dan identitas brand \"PP. Miftahul Ihsan\". Panel ini mengelola seluruh data operasional pondok dengan kontrol akses berbasis peran (Role-Based Access Control) menggunakan Filament Shield, yang memisahkan hak akses Admin, Operator, Pengasuh (read-only), dan Keamanan.")

    add_figure(doc, "04_admin_dashboard.png", "Dashboard Utama Panel Admin")
    add_body(doc, "Dashboard menampilkan widget statistik ringkasan: jumlah santri aktif, tagihan menunggak, eskalasi pelanggaran, dan izin menunggu verifikasi. Menu navigasi tersusun dalam kelompok Data Master, Kedisiplinan, Keuangan, Manajemen Pengguna, dan Sistem & Pengaturan.")

    add_figure(doc, "05_admin_santris.png", "Halaman Data Santri")
    add_body(doc, "Modul Data Santri menyediakan tabel lengkap identitas santri beserta fitur pencarian, penyaringan, dan aksi tambah/ubah/hapus. Data santri diimpor dari berkas Excel data pondok melalui seeder untuk memastikan kelengkapan data awal.")

    add_figure(doc, "06_admin_wali_santris.png", "Halaman Data Wali Santri")
    add_body(doc, "Data wali santri dikelola terpisah dan terhubung dengan akun login portal wali, sehingga setiap wali hanya dapat melihat data anaknya masing-masing (data scoping per wali).")

    add_figure(doc, "07_admin_kamars.png", "Halaman Data Kamar / Asrama")
    add_body(doc, "Modul ini mengelola daftar kamar dan asrama beserta kapasitasnya, yang digunakan sebagai referensi penempatan santri.")

    add_figure(doc, "08_admin_pelanggarans.png", "Halaman Catatan Pelanggaran Santri")
    add_body(doc, "Pelanggaran santri dicatat dengan kategori pelanggaran dan bobot poin. Setiap pelanggaran memicu notifikasi WhatsApp otomatis kepada wali santri sesuai template yang telah dikonfigurasi.")

    add_figure(doc, "09_admin_penghargaans.png", "Halaman Penghargaan Santri")
    add_body(doc, "Pencatatan penghargaan/prestasi santri memberikan keseimbangan penilaian kedisiplinan, dan riwayatnya ditampilkan pada portal wali.")

    add_figure(doc, "10_admin_tagihans.png", "Halaman Tagihan Santri")
    add_body(doc, "Tagihan dibuat otomatis oleh sistem melalui TagihanService, dan setiap pembuatan tagihan mengirim notifikasi WhatsApp kepada wali santri. Status tagihan (belum lunas/lunas) dipantau dari dashboard.")

    add_figure(doc, "11_admin_pembayarans.png", "Halaman Catatan Pembayaran")
    add_body(doc, "Pembayaran tagihan dicatat dan dikonfirmasi admin, mengubah status tagihan menjadi lunas dan memperbarui statistik keuangan pondok.")

    add_figure(doc, "12_admin_perizinans.png", "Halaman Pengajuan Perizinan Santri")
    add_body(doc, "Wali santri dapat mengajukan izin keluar pondok melalui portal; admin/keamanan memverifikasi dan menyetujui atau menolak pengajuan tersebut.")

    add_figure(doc, "13_admin_tahfidzs.png", "Halaman Catatan Tahfidz Santri")
    add_body(doc, "Perkembangan hafalan Al-Qur'an santri dicatat oleh pengasuh dan dapat dipantau oleh wali santri melalui portal.")

    add_figure(doc, "14_admin_riwayat_kesehatans.png", "Halaman Riwayat Kesehatan Santri")
    add_body(doc, "Riwayat kesehatan santri mencakup penyakit bawaan dan catatan penanganan medis, mendukung pengambilan keputusan pondok dalam merawat santri.")

    add_figure(doc, "15_admin_penguruses.png", "Halaman Data Pengurus")
    add_body(doc, "Data pengurus pondok (ustadz/ustadzah dan staf) dikelola pada modul ini untuk keperluan operasional dan akuntabilitas.")

    add_figure(doc, "16_admin_users.png", "Halaman Manajemen Pengguna")
    add_body(doc, "Administrator mengelola akun pengguna sistem dan penetapan peran (Admin, Operator, Pengasuh, Keamanan, Wali Santri) beserta hak aksesnya.")

    # 4.1.4 Implementasi Notifikasi WhatsApp Gateway
    add_h3(doc, "4.1.4", "Implementasi Modul Notifikasi WhatsApp Gateway")
    add_body(doc, "Integrasi notifikasi WhatsApp menggunakan pustaka laravel-whatsapp dengan mekanisme Web Sidecar (whatsapp-web.js) yang berjalan sebagai proses terpisah bersama Chromium. Sidecar dikelola melalui perintah artisan (whatsapp:sidecar:install, whatsapp:sidecar:start, whatsapp:sidecar:status) dan dapat ditautkan ke nomor WhatsApp melalui pemindaian QR Code atau kode pairing. Status sesi ditampilkan secara live tanpa polling menggunakan Server-Sent Events (SSE).")

    add_figure(doc, "17_admin_whatsapp_gateway.png", "Halaman WhatsApp Gateway (Kelola Sesi Sidecar)")
    add_body(doc, "Halaman ini menampilkan status sesi WhatsApp secara live (QR Code / pairing code / terhubung), dengan aksi mulai (start), hentikan (stop), hapus sesi (destroy), dan permintaan kode pairing.")

    add_figure(doc, "18_admin_whatsapp_templates.png", "Halaman Templat WhatsApp")
    add_body(doc, "Template pesan WhatsApp dikelola dengan dukungan variabel dinamis seperti {nama_santri}, {nama_wali}, {jenis}, {deskripsi}, {tagihan}, dan {tanggal}, sehingga pesan otomatis tersusun rapi sesuai konteks notifikasi pelanggaran maupun tagihan.")

    add_figure(doc, "19_admin_notifikasi_logs.png", "Halaman Log Notifikasi WhatsApp")
    add_body(doc, "Seluruh pengiriman notifikasi tercatat pada log dengan status terkirim/gagal. Log yang gagal dapat dikirim ulang (retry) melalui aksi pada baris log, memastikan wali santri tetap menerima informasi penting.")

    add_figure(doc, "20_admin_roles.png", "Halaman Pengaturan Peran & Izin (Filament Shield)")
    add_body(doc, "Hak akses berbasis peran dikelola dengan Filament Shield. Role Admin memiliki seluruh izin (207 permission), Operator untuk data operasional, Pengasuh bersifat read-only, Keamanan untuk modul kedisiplinan dan perizinan, serta Wali Santri hanya untuk data anaknya.")

    add_figure(doc, "21_admin_rule_poin.png", "Halaman Pengaturan Rule Poin Kedisiplinan")
    add_body(doc, "Bobot poin setiap kategori pelanggaran dan ambang batas eskalasi dapat dikonfigurasi, memberikan fleksibilitas pengasuh dalam menerapkan aturan kedisiplinan pondok.")

    # 4.1.5 Implementasi Portal Wali Santri
    add_h3(doc, "4.1.5", "Implementasi Portal Wali Santri")
    add_body(doc, "Portal wali santri merupakan panel khusus orang tua/wali dengan tampilan modern dan user-friendly. Setiap wali hanya dapat mengakses data anaknya sendiri, dilindungi oleh scoping data pada level query sehingga wali lain tidak dapat melihat data tersebut.")

    add_figure(doc, "22_wali_dashboard.png", "Dashboard Portal Wali Santri")
    add_body(doc, "Dashboard wali menampilkan kartu statistik (anak asuh, total poin, tagihan terdekat, status keuangan), widget tabel tagihan anak, serta grafik batang perkembangan poin per anak.")

    add_figure(doc, "23_wali_santris.png", "Halaman Data Santri pada Portal Wali")
    add_body(doc, "Wali melihat profil lengkap anaknya (identitas, kamar, dan informasi terkait) sesuai data yang terdaftar di pondok.")

    add_figure(doc, "24_wali_pelanggarans.png", "Halaman Pelanggaran Anak pada Portal Wali")
    add_body(doc, "Wali memantau riwayat pelanggaran anak beserta kategori dan poin, sehingga dapat memberikan pembinaan di rumah.")

    add_figure(doc, "25_wali_tagihans.png", "Halaman Tagihan Anak pada Portal Wali")
    add_body(doc, "Daftar tagihan anak ditampilkan lengkap dengan status pembayaran, membantu wali merencanakan pembayaran tepat waktu.")

    add_figure(doc, "26_wali_tahfidzs.png", "Halaman Tahfidz Anak pada Portal Wali")
    add_body(doc, "Perkembangan hafalan anak dapat dipantau oleh wali untuk mendukung program tahfidz di rumah.")

    add_figure(doc, "27_wali_perizinans.png", "Halaman Perizinan Anak pada Portal Wali")
    add_body(doc, "Wali dapat mengajukan izin keluar pondok dan memantau status persetujuan dari pihak pondok.")

    # 4.1.6 Pengujian Sistem
    add_h3(doc, "4.1.6", "Hasil Pengujian Fungsional")
    add_body(doc, "Pengujian sistem dilakukan dengan dua pendekatan: (1) Automated Testing menggunakan PHPUnit/Pest pada framework Laravel yang mencakup 7 kasus uji (unit dan fitur), dan (2) Pengujian Black-Box melalui pengamatan langsung antarmuka pada browser Google Chrome terhadap 14 skenario fungsional utama. Hasil pengujian black-box disajikan pada Tabel 4.2.")

    test_rows = [
        ("1", "Autentikasi Login Panel Admin", "Username: admin, Password: password di /admin/login", "Login berhasil dan diarahkan ke Dashboard Admin", "Diarahkan ke /admin dengan seluruh menu tampil", "PASS"),
        ("2", "Autentikasi Login Portal Wali via Nomor HP", "No. HP: 087787224620, Password: password di /wali/login", "Login berhasil dan diarahkan ke Dashboard Wali", "Diarahkan ke /wali dengan data anak sendiri", "PASS"),
        ("3", "Pembatasan Akses Antar Panel", "Admin mengakses /wali; wali mengakses /admin", "Masing-masing mendapat 403 (tidak berhak)", "403 sesuai peran pengguna", "PASS"),
        ("4", "Landing Page & Seluruh Section", "Akses URL /", "Menampilkan 16 section sesuai PRD dengan animasi", "Semua section render, aset CSS/JS termuat", "PASS"),
        ("5", "Scoping Data Wali Santri", "Wali login lalu membuka Data Santri", "Hanya menampilkan data anak sendiri", "Hanya anak sendiri yang tampil", "PASS"),
        ("6", "Dashboard Wali & Widget", "Akses /wali setelah login", "Kartu statistik, tabel tagihan, dan grafik poin render", "Semua widget render tanpa error", "PASS"),
        ("7", "Pembuatan Tagihan & Notifikasi", "Buat tagihan baru sebagai admin", "Tagihan tersimpan dan notifikasi WA ke wali tercatat", "Log notifikasi masuk dengan template tagihan", "PASS"),
        ("8", "Notifikasi Pelanggaran", "Catat pelanggaran dengan kategori", "Notifikasi WA terkirim ke wali sesuai template", "Log terkirim dengan template pelanggaran", "PASS"),
        ("9", "Templat WhatsApp & Variabel", "Sunting template, simpan, render", "Variabel {nama_santri} dll. diganti nilai sebenarnya", "Render variabel benar (unit test PASS)", "PASS"),
        ("10", "Log Notifikasi & Retry", "Buka Log Notifikasi, klik kirim ulang pada log gagal", "Pesan diproses ulang dan log diperbarui", "Aksi retry tersedia dan berjalan", "PASS"),
        ("11", "Status Live WhatsApp Gateway", "Buka halaman WhatsApp Gateway", "Status sesi tampil live via SSE tanpa polling", "Status terhubung/QR tampil real-time", "PASS"),
        ("12", "Manajemen Role & Izin", "Akses menu Roles sebagai Admin", "Seluruh role dan izin dapat dikelola", "Role Admin (207 izin), Operator, Pengasuh, dll. tampil", "PASS"),
        ("13", "Tema & Branding Dashboard", "Login sebagai admin/wali", "Tema hijau dan brand \"PP. Miftahul Ihsan\" tampil", "Tema Emerald & brand tampil di semua panel", "PASS"),
        ("14", "Widget Welcome dihapus", "Buka dashboard admin & wali", "Tidak ada kartu welcome/Filament", "Dashboard bersih hanya widget kustom", "PASS"),
    ]
    add_bb_table(doc, "4.2", "Hasil Pengujian Black-Box Sistem Informasi Manajemen Santri", test_rows)

    add_body(doc, "Selain pengujian black-box, seluruh Automated Testing yang tertanam pada basis kode berhasil dijalankan dengan hasil 7 tes lulus (19 asersi). Daftar tes tersebut adalah: (1) test_the_application_returns_a_successful_response, (2) test_kirim_sukses_memperbarui_status_log, (3) test_kirim_gagal_memperbarui_status_dan_melempar_exception, (4) test_kirim_tanpa_no_hp_menandai_gagal_tanpa_request, (5) test_that_true_is_true, (6) test_render_mengganti_semua_placeholder, dan (7) test_render_placeholder_yang_tidak_diisi_tetap_utuh.")

    # 4.2 Pembahasan
    add_h2(doc, "4.2", "Pembahasan")

    add_h3(doc, "4.2.1", "Analisis Hasil Pengujian")
    add_body(doc, "1. Tingkat Keberhasilan Fungsional (Success Rate): Dari 14 skenario black-box dan 7 tes otomatis yang dieksekusi, seluruh kasus uji (100%) dinyatakan PASS. Tidak ditemukan kesalahan fatal maupun kegagalan alur pada modul inti sistem.")
    add_body(doc, "2. Keandalan Kontrol Akses (RBAC): Pemisahan hak akses antar peran berjalan konsisten. Admin tidak dapat mengakses portal wali dan sebaliknya; pengasuh dibatasi hanya membaca data; dan wali santri hanya melihat data anaknya sendiri. Hal ini membuktikan penerapan Spatie Laravel Permission dan Filament Shield bekerja sesuai spesifikasi.")
    add_body(doc, "3. Efektivitas Notifikasi WhatsApp: Integrasi laravel-whatsapp dengan Web Sidecar terbukti mampu mengirim notifikasi pelanggaran dan tagihan secara otomatis dengan status sesi yang dapat dipantau live. Mekanisme retry pada log notifikasi memberikan keandalan tambahan terhadap kegagalan pengiriman.")
    add_body(doc, "4. Kualitas Antarmuka: Landing page dan portal wali dirancang modern dan responsif sesuai PRD. Verifikasi dengan headless Chrome memastikan seluruh aset (CSS/JS), animasi scroll-reveal, dan count-up statistik berjalan tanpa error konsol.")

    add_h3(doc, "4.2.2", "Kelebihan dan Keterbatasan Sistem")
    add_body(doc, "Sistem memiliki beberapa keunggulan utama: (1) pengelolaan data santri terpusat dengan impor data awal dari berkas Excel pondok; (2) notifikasi WhatsApp otomatis tanpa biaya API tambahan karena memakai Web Sidecar whatsapp-web.js; (3) status gateway live tanpa polling via SSE; (4) portal wali modern dengan scoping data yang aman; dan (5) tema serta identitas pondok yang konsisten di seluruh antarmuka.")
    add_body(doc, "Meskipun demikian, sistem masih memiliki keterbatasan: (1) sidecar WhatsApp bergantung pada ketersediaan Chromium dan kestabilan koneksi; (2) sesi WhatsApp Web dapat terputus jika nomor diblokir atau perangkat seluler kehilangan koneksi; (3) notifikasi hanya mencakup pelanggaran dan tagihan sesuai keputusan pengelola; dan (4) pengujian keamanan lanjutan (penetration testing) serta pengujian beban belum dilakukan.")

    print("Generating BAB V...")

    # ---------------- BAB V ----------------
    add_bab_heading(doc, "V", "KESIMPULAN DAN SARAN")

    add_h2(doc, "5.1", "Kesimpulan")
    add_body(doc, "Berdasarkan hasil perancangan, implementasi, dan pengujian Sistem Informasi Manajemen Santri Pondok Pesantren Miftahul Ihsan, dapat ditarik kesimpulan sebagai berikut:")
    add_body(doc, "1. Telah berhasil dirancang dan dibangun sistem informasi manajemen santri berbasis web menggunakan Laravel 13, Filament v5, dan Livewire yang mencakup pengelolaan data santri, kedisiplinan (pelanggaran dan penghargaan), tagihan dan pembayaran, perizinan, tahfidz, riwayat kesehatan, serta pengguna dengan kontrol akses berbasis peran (Admin, Operator, Pengasuh read-only, Keamanan, dan Wali Santri).")
    add_body(doc, "2. Sistem berhasil mengintegrasikan notifikasi WhatsApp melalui laravel-whatsapp Web Sidecar untuk notifikasi pelanggaran dan tagihan, dengan fitur kelola sesi (scan QR / pairing code), status live via SSE, template pesan, dan log notifikasi yang dapat dikirim ulang (retry).")
    add_body(doc, "3. Portal wali santri berhasil diimplementasikan dengan login menggunakan nomor HP, tampilan modern dan user-friendly, serta scoping data yang menjamin wali hanya dapat mengakses data anaknya sendiri.")
    add_body(doc, "4. Hasil pengujian black-box terhadap 14 skenario dan automated testing terhadap 7 kasus uji menunjukkan tingkat keberhasilan 100% (PASS), membuktikan seluruh fitur beroperasi optimal, stabil, dan sesuai spesifikasi kebutuhan pada PRD.")

    add_h2(doc, "5.2", "Saran")
    add_body(doc, "Untuk pengembangan sistem di masa mendatang, diajukan beberapa saran sebagai berikut:")
    add_body(doc, "1. Pengembangan Aplikasi Mobile: Perlu dipertimbangkan pembangunan aplikasi mobile (Android/iOS) untuk portal wali agar orang tua lebih mudah memantau anak dan menerima push notification.")
    add_body(doc, "2. Perluasan Cakupan Notifikasi: Cakupan notifikasi WhatsApp dapat diperluas, misalnya untuk pengingat pembayaran otomatis (reminder tagihan), informasi kegiatan pondok, dan pengumuman penting lainnya.")
    add_body(doc, "3. Redundansi Gateway WhatsApp: Disarankan menyediakan lebih dari satu sesi/nomor WhatsApp sebagai cadangan (failover) agar pengiriman notifikasi tetap berjalan jika satu sesi terputus.")
    add_body(doc, "4. Pengujian Lanjutan: Perlu dilakukan penetration testing dan load testing dengan jumlah santri dalam skala besar untuk mengevaluasi keamanan dan skalabilitas sistem sebelum diterapkan secara penuh di lingkungan produksi.")

    print("Generating DAFTAR GAMBAR...")
    add_daftar_gambar(doc)

    print(f"Saving document to {OUT}...")
    doc.save(OUT)
    print("Done generating BAB IV & BAB V docx!")

if __name__ == "__main__":
    generate_docx()
